#!/usr/bin/env python3
"""
CChO Acid Ranking Question Generator

Generates Chinese Chemistry Olympiad style acidity ranking questions
with ACS 1996 standard chemical structures embedded as CDXML in DOCX.

Usage: python generate.py [--count N] [--output path.docx]
"""

import os, re, io, sys, json, random, hashlib, struct, uuid, argparse, datetime
from pathlib import Path
from collections import defaultdict

# ── Dependencies ──────────────────────────────────────────────────
try:
    from rdkit import Chem, RDLogger
    from rdkit.Chem import AllChem, Draw, rdMolDescriptors
    from rdkit.Chem.Draw import rdMolDraw2D
    RDLogger.logger().setLevel(RDLogger.ERROR)  # Suppress warnings
    HAS_RDKIT = True
except ImportError:
    HAS_RDKIT = False
    print("WARNING: rdkit not installed. Install with: pip install rdkit")

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("WARNING: python-docx not installed. Install with: pip install python-docx")

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import olefile
    HAS_OLE = True
except ImportError:
    HAS_OLE = False
    print("INFO: olefile not available, OLE embedding will use simple wrapper")

# ── Constants ──────────────────────────────────────────────────────
# ACS 1996 standard settings
ACS_BOND_LENGTH = 30.0      # RDKit fixedBondLength
ACS_LINE_WIDTH = 2.5        # RDKit bondLineWidth multiplier
ACS_FONT_SIZE = 24          # RDKit atomLabelFontSize (10pt equivalent)
ACS_FONT = "Arial"
ACS_IMAGE_WIDTH = 400
ACS_IMAGE_HEIGHT = 300

# CDXML scale: RDKit angstroms → CDXML units (approximate)
CDXML_SCALE = 28.35  # points per cm equivalent

# ── Compound Database ──────────────────────────────────────────────

class Compound:
    """A single compound entry from the pKa database."""
    def __init__(self, smiles, name, pka_h2o=None, pka_dmso=None, category=None, pka_type=None):
        self.smiles = smiles
        self.name = name
        self.pka_h2o = pka_h2o      # float or None
        self.pka_dmso = pka_dmso    # float or None
        self.category = category
        self.pka_type = pka_type

    def best_pka(self):
        """Return the best available pKa value (prefer H2O)."""
        if self.pka_h2o is not None:
            return self.pka_h2o
        return self.pka_dmso

    def __repr__(self):
        return f"Compound({self.name}, pKa={self.best_pka()})"


def parse_acid_library(md_path):
    """Parse 酸性数据.md into a list of Compound objects.
    Also loads IUPAC high-confidence data from the cloned repo.
    Returns (compounds_list, stats_dict)."""
    compounds = []

    # ── Load 酸性数据.md ──
    skill_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    if not md_path:
        md_path = os.path.join(skill_dir, "酸性数据.md")
    if not os.path.exists(md_path):
        alt = os.path.join(os.path.expanduser("~"), "Downloads", "酸性数据.md")
        if os.path.exists(alt):
            md_path = alt

    if os.path.exists(md_path):
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        compounds.extend(_parse_md_tables(content))
        print(f"  [酸性数据.md] {len(compounds)} compounds loaded")
    else:
        print(f"  [酸性数据.md] NOT FOUND at {md_path}")

    # ── Load IUPAC data ──
    iupac_csv = os.path.join(skill_dir, "iupac_high-confidence_v2_3.csv")
    iupac_count = 0
    if os.path.exists(iupac_csv):
        try:
            from parse_iupac import get_iupac_compounds
            iupac_compounds = get_iupac_compounds(iupac_csv)
            compounds.extend(iupac_compounds)
            iupac_count = len(iupac_compounds)
            print(f"  [IUPAC] {iupac_count} compounds loaded")
        except ImportError:
            print("  [IUPAC] parse_iupac module not found, skipping")
    else:
        print(f"  [IUPAC] CSV not found at {iupac_csv}")

    return compounds, {'md': len(compounds) - iupac_count, 'iupac': iupac_count}


def _parse_md_tables(content):
    """Parse markdown tables from 酸性数据.md into Compound list."""
    compounds = []
    current_category = ""
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('## Page'):
            pass
        elif line.startswith('### ') and not line.startswith('####'):
            current_category = line.replace('### ', '').strip()

        # Parse table rows
        # Match: | `SMILES` | Name | pKa | pKa(DMSO) |
        table_row = re.match(
            r'^\|\s*`?([^`|]+?)`?\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.*?)\s*\|$',
            line
        )
        if table_row:
            smiles_raw = table_row.group(1).strip()
            name_raw = table_row.group(2).strip()
            pka_raw = table_row.group(3).strip()
            dmso_raw = table_row.group(4).strip()

            # Skip header rows
            if smiles_raw in ('SMILES', '---', ':---') or '---' in smiles_raw:
                i += 1
                continue
            if 'Compound Name' in name_raw or 'pKa' in name_raw:
                i += 1
                continue

            # Clean SMILES
            smiles = smiles_raw.strip().strip('`').strip()
            # Remove backtick formatting
            smiles = re.sub(r'`', '', smiles)

            # Parse pKa
            pka_h2o = None
            pka_dmso = None

            # Handle multiple pKa values (e.g., "4.76" or "1.9, 7.21" or "19-20")
            if pka_raw and pka_raw.strip() and pka_raw.strip() != '–' and pka_raw.strip() != '-':
                pka_str = pka_raw.strip()
                # Take the first number (lowest pKa = strongest acid)
                nums = re.findall(r'[-]?\d+\.?\d*', pka_str)
                if nums:
                    try:
                        pka_h2o = float(nums[0])
                    except ValueError:
                        pass

            if dmso_raw and dmso_raw.strip() and dmso_raw.strip() != '–':
                nums = re.findall(r'[-]?\d+\.?\d*', dmso_raw.strip())
                if nums:
                    try:
                        pka_dmso = float(nums[0])
                    except ValueError:
                        pass

            if smiles and len(smiles) > 1 and smiles not in ('', ' ', '-'):
                # Clean up non-standard SMILES
                cleaned = clean_smiles(smiles)
                # Basic SMILES validation
                if Chem.MolFromSmiles(cleaned) is not None or is_inorganic_smiles(cleaned):
                    smiles = cleaned
                    compounds.append(Compound(
                        smiles=smiles,
                        name=name_raw.strip(),
                        pka_h2o=pka_h2o,
                        pka_dmso=pka_dmso,
                        category=current_category
                    ))

        i += 1

    return compounds


def is_inorganic_smiles(smiles):
    """Check if a SMILES represents a simple inorganic species."""
    inorganic = {'O', 'S', 'Br', 'Cl', 'F', 'OO', 'OCl', 'C#N', 'SC#N', 'N'}
    return smiles in inorganic


def clean_smiles(smiles):
    """Clean up non-standard SMILES abbreviations that RDKit may not parse."""
    # Replace common non-standard patterns
    smiles = smiles.replace('PhS(=O)(=O)CPh', 'O=S(=O)(Cc1ccccc1)c2ccccc2')
    smiles = smiles.replace('PhS(=O)Cc1ccccc1', 'O=S(Cc1ccccc1)c1ccccc1')
    smiles = smiles.replace('PhSCc1ccccc1', 'S(Cc1ccccc1)c1ccccc1')
    smiles = smiles.replace('CS(=NTs)C', 'CS(=NS(=O)(=O)c1ccc(C)cc1)C')
    smiles = smiles.replace('CS(=O)(=NTs)C', 'CS(=O)(=NS(=O)(=O)c1ccc(C)cc1)C')
    smiles = smiles.replace('PhP+(C)Ph', 'C[P+](c1ccccc1)c1ccccc1')
    # General "Ph" → "c1ccccc1" for single Ph
    # But careful: don't break [N+] patterns etc
    if 'Ph' in smiles and smiles not in ['Ph', '[PH4+]']:
        smiles = re.sub(r'(?<!\[)Ph(?!\+)', 'c1ccccc1', smiles)
    # Fix Ts (tosyl) → appropriate SMILES
    if 'Ts' in smiles and 'NTs' not in smiles:
        smiles = smiles.replace('Ts', 'S(=O)(=O)c1ccc(C)cc1')
    return smiles


# ── Compound Tagging ──────────────────────────────────────────────

def tag_compound(comp):
    """Tag a compound with its testing theme.

    Returns (theme_type, family_key, difficulty)
    theme_type: 'inorganic' | 'substituent' | 'functional_group' | 'exception'
    family_key: identifier for grouping (e.g. 'benzoic_acid', 'phenol', 'acetic_acid')
    difficulty: 1 (easy) | 2 (medium) | 3 (hard)
    """
    name = comp.name.lower()
    smiles = comp.smiles
    pka = comp.best_pka()

    # ── Unprotonated pKaH1 entries (conjugate acids whose SMILES
    #     still represents the neutral base) → isolate them ──
    if comp.pka_type and comp.pka_type.startswith('pKaH') and '+' not in smiles:
        return ('conjugate_acid', 'conjugate_acid_neutral', 3)

    # ── Inorganic ──
    inorganic_names = {
        'h₂o', 'h₃o⁺', 'h₂s', 'hbr', 'hcl', 'hf', 'hocl', 'hclo₄', 'hcn',
        'hn₃', 'hscn', 'h₂so₃', 'h₂so₄', 'h₃po₄', 'hno₃', 'hno₂',
        'h₂cro₄', 'ch₃so₃h', 'cf₃so₃h', 'nh₄cl', 'b(oh)₃', 'hooh',
        'h₂o₂', 'water', 'methanesulfonic', 'triflic', 'perchloric',
        'boric', 'chromic', 'hydrazoic', 'thiocyanic', 'hypochlorous',
        'hydrogen sulfide', 'hydrogen bromide', 'hydrogen chloride',
        'hydrogen fluoride', 'hydrogen cyanide', 'nitric', 'nitrous',
        'sulfuric', 'sulfurous', 'phosphoric', 'ammonium',
    }
    for kw in inorganic_names:
        if kw in name or kw in smiles.lower():
            return ('inorganic', 'inorganic', 1)

    # Check SMILES for simple inorganic patterns
    simple = {'O', 'S', 'Br', 'Cl', 'F', 'OO', 'OCl', 'C#N', 'SC#N', 'N'}
    if smiles.strip() in simple:
        return ('inorganic', 'inorganic', 1)
    if '[OH3+]' in smiles or 'O=S(=O)(O)C(F)(F)F' in smiles:
        return ('inorganic', 'inorganic', 1)

    # ── Famous Exceptions ──
    exceptions = [
        ('meldrum', 'meldrums_acid', 3),
        ('dimedone', 'dimedone', 3),
        ('barbituric', 'barbituric_acid', 3),
        ('squaric', 'squaric_acid', 3),
        ('acetylacetone', 'acetylacetone', 2),
        ('malononitrile', 'malononitrile', 3),
        ('nitromethane', 'nitromethane', 2),
        ('triflone', 'triflone', 3),
        ('fulvene', 'fulvene', 3),
        ('cyclopentadiene', 'cyclopentadiene', 2),
        ('fluorene', 'fluorene', 2),
        ('indene', 'indene', 2),
        ('pentane-2,4-dione', 'acetylacetone', 2),
        ('propanedinitrile', 'malononitrile', 3),
        ('diketene', 'diketene', 3),
        ('cyanodinitromethane', 'cyanodinitromethane', 3),
        ('bis(ethylsulfonyl)methane', 'disulfone', 3),
        ('bis(phenylsulfonyl)methane', 'disulfone', 3),
        ('Meldrum', 'meldrums_acid', 3),
        ('2,2-dimethyl-1,3-dioxane-4,6-dione', 'meldrums_acid', 3),
    ]
    for kw, family, diff in exceptions:
        if kw in name:
            return ('exception', family, diff)

    # ── Substituent Effect Families ──
    # Benzoic acids (X-C6H4COOH)
    if 'benzoic' in name or ('c1ccccc1' in smiles and 'C(=O)O' in smiles):
        if 'nitro' in name:
            return ('substituent', 'benzoic_acid', 1)
        if 'chloro' in name or 'bromo' in name or 'methoxy' in name:
            return ('substituent', 'benzoic_acid', 1)
        return ('substituent', 'benzoic_acid', 1)

    # Phenols (X-C6H4OH)
    if 'phenol' in name or 'naphthol' in name:
        return ('substituent', 'phenol', 1)

    # Haloacetic acids (X-CH2COOH)
    if 'acetic' in name or 'chloroacetic' in name or 'bromoacetic' in name or \
       'iodoacetic' in name or 'fluoroacetic' in name or 'dichloroacetic' in name or \
       'trichloroacetic' in name or 'trifluoroacetic' in name:
        return ('substituent', 'acetic_acid', 1)

    # Alcohols
    if 'methanol' in name or 'ethanol' in name or 'propanol' in name or \
       'butanol' in name or 'hexanol' in name or 'trifluoroethanol' in name:
        return ('substituent', 'alcohol', 2)

    # Amines / ammonium
    if 'ammonium' in name or 'anilinium' in name:
        return ('substituent', 'ammonium', 2)
    if 'amine' in name or 'aniline' in name:
        return ('substituent', 'amine', 2)

    # C-H acids with substituent patterns
    if 'ketone' in name or 'acetone' in name or 'acetophenone' in name:
        return ('substituent', 'ketone', 2)

    # ── Functional Group Comparison ──
    # Carboxylic acids (general) — but NOT already tagged as substituent
    is_subst_acids = {'acetic', 'chloroacetic', 'bromoacetic', 'iodoacetic',
                      'fluoroacetic', 'dichloroacetic', 'trichloroacetic', 'trifluoroacetic'}
    name_lower = name
    if ('carboxylic' in name_lower or 'formic' in name_lower or 'acrylic' in name_lower or
        'oxalic' in name_lower or 'peracetic' in name_lower):
        return ('functional_group', 'carboxylic_acid', 2)

    # Amides
    if 'amide' in name or 'carbamate' in name or 'urea' in name:
        return ('functional_group', 'amide', 2)

    # Imides, sulfonamides
    if 'imide' in name or 'sulfonamide' in name or 'phthalimide' in name:
        return ('functional_group', 'imide', 2)

    # Heterocyclic NH acids
    if 'pyrrole' in name or 'imidazole' in name or 'triazole' in name or \
       'tetrazole' in name or 'indole' in name or 'carbazole' in name or \
       'benzimidazole' in name or 'pyrazole' in name:
        return ('functional_group', 'heterocycle_nh', 2)

    # Sulfur acids (sulfides, sulfoxides, sulfones, C-H adjacent to S)
    if 'sulfinic' in name or 'sulfonic' in name or 'sulfoxide' in name or \
       'sulfone' in name or 'sulfide' in name or 'thiol' in name or \
       'mercapto' in name or 'sulfonium' in name or 'sulfilimine' in name:
        return ('functional_group', 'sulfur_acid', 2)

    # Nitriles, nitroalkanes
    if 'nitrile' in name or 'cyano' in name or 'cyanide' in name:
        return ('functional_group', 'nitrile', 2)
    if 'nitro' in name and ('methane' in name or 'ethane' in name or 'propane' in name):
        return ('functional_group', 'nitroalkane', 2)

    # Esters
    if 'ester' in name or 'EtO' in name or 't-BuO' in name:
        return ('functional_group', 'ester', 3)

    # Ketones
    if 'ketone' in name or 'acetophenone' in name or 'cyclohexanone' in name or \
       'diketone' in name:
        return ('functional_group', 'ketone', 2)

    # Hydrocarbons (very weak C-H acids)
    if 'methane' in name or 'ethane' in name or 'ethylene' in name or \
       'acetylene' in name or 'toluene' in name or \
       'propene' in name or 'pentyne' in name or 'diphenylmethane' in name or \
       'triphenylmethane' in name or 'fluorene' in name or 'indene' in name:
        return ('functional_group', 'hydrocarbon', 3)

    # Ethers
    if 'ether' in name or 'CH₃OPh' in name or 'MeO' in name or 'PhO' in name:
        return ('functional_group', 'ether', 3)

    # Phosphorus compounds
    if 'phosph' in name:
        return ('functional_group', 'phosphorus', 3)

    # Selenium
    if 'selen' in name:
        return ('functional_group', 'selenium', 3)

    # Oximes, hydrazones
    if 'oxime' in name or 'hydrazone' in name or 'hydrazide' in name:
        return ('functional_group', 'oxime_hydrazone', 2)

    # Amidines, guanidines
    if 'amidine' in name or 'guanidine' in name:
        return ('functional_group', 'amidine', 2)

    # Protonated species
    if 'protonated' in name or 'onium' in name:
        return ('functional_group', 'protonated', 1)

    # ── Default ──
    return ('functional_group', 'other', 3)


# ── Question Generation ────────────────────────────────────────────

def select_compounds_for_ranking(compounds, n=5, pka_range=None):
    """
    Select n compounds suitable for ranking.
    Compounds should have similar but distinguishable pKa values.
    """
    # Filter compounds with known pKa
    valid = [c for c in compounds if c.best_pka() is not None]

    if pka_range:
        valid = [c for c in valid if pka_range[0] <= c.best_pka() <= pka_range[1]]

    if len(valid) < n:
        return random.sample(valid, min(len(valid), n)) if valid else []

    # Try to find compounds with reasonable pKa spread (2-8 units)
    random.shuffle(valid)

    best_group = None
    best_spread = -1
    for _ in range(min(50, len(valid))):
        group = random.sample(valid, n)
        pkas = [c.best_pka() for c in group]
        spread = max(pkas) - min(pkas)
        if 3 <= spread <= 12:
            best_group = group
            break
        if spread > best_spread:
            best_spread = spread
            best_group = group

    return best_group if best_group else valid[:n]


def _build_theme_groups(compounds):
    """Build compound pools grouped by testing theme."""
    theme_pools = {
        'inorganic': [],
        'substituent': defaultdict(list),   # family_key → [compounds]
        'functional_group': defaultdict(list),
        'exception': defaultdict(list),
        'conjugate_acid': defaultdict(list),
    }

    for c in compounds:
        if c.best_pka() is None:
            continue
        theme_type, family, difficulty = tag_compound(c)
        # Attach difficulty info
        c.difficulty = difficulty
        c.family = family
        c.theme_type = theme_type

        if theme_type == 'inorganic':
            theme_pools['inorganic'].append(c)
        elif theme_type == 'substituent':
            theme_pools['substituent'][family].append(c)
        elif theme_type == 'functional_group':
            theme_pools['functional_group'][family].append(c)
        elif theme_type == 'exception':
            theme_pools['exception'][family].append(c)
        elif theme_type == 'conjugate_acid':
            theme_pools['conjugate_acid'][family].append(c)

    return theme_pools


def _pick_substituent_question(theme_pools, n_compounds=4, difficulty=1):
    """Pick a substituent-effect question."""
    # Find families with enough members
    candidates = []
    for family, members in theme_pools['substituent'].items():
        if len(members) >= n_compounds:
            # Filter by difficulty
            matched = [c for c in members if c.difficulty <= difficulty]
            if len(matched) >= n_compounds:
                candidates.append((family, matched))
            elif len(members) >= n_compounds:
                candidates.append((family, members))

    if not candidates:
        return None

    family, pool = random.choice(candidates)
    selected = random.sample(pool, min(n_compounds, len(pool)))
    selected.sort(key=lambda c: c.best_pka())

    # Optionally sprinkle an inorganic
    inorganics = theme_pools['inorganic']
    if inorganics and random.random() < 0.4 and len(selected) < n_compounds + 1:
        ino = random.choice(inorganics)
        if ino not in selected:
            selected.append(ino)

    selected.sort(key=lambda c: c.best_pka())
    return selected


def _validate_no_cross(compounds):
    """Ensure compounds sorted by pKa form contiguous theme blocks.
    
    OK:     [substituent, substituent] > [functional_group, functional_group]
    NOT OK: [substituent, functional_group, substituent, functional_group]
    
    Returns True if themes are grouped (no interleaving).
    """
    sorted_compounds = sorted(compounds, key=lambda c: c.best_pka() if c.best_pka() is not None else 999)
    
    seen_families = []  # ordered list of families as they appear
    for c in sorted_compounds:
        family = getattr(c, 'family', tag_compound(c)[1])
        if not seen_families or seen_families[-1] != family:
            if family in seen_families:
                # This family appeared earlier, then a different family, now reappears → crossing!
                return False
            seen_families.append(family)
    return True


def _pick_functional_group_question(theme_pools, n_compounds=4, difficulty=2):
    """Pick a functional-group comparison question where themes don't interleave."""
    pools = theme_pools['functional_group']

    eligible = {}
    for family, members in pools.items():
        filtered = [c for c in members if c.difficulty <= difficulty]
        if len(filtered) >= 1:
            eligible[family] = filtered
        elif len(members) >= 1:
            eligible[family] = members

    if len(eligible) < 2:
        return None

    # Try up to 20 times to find a non-crossing group
    for _ in range(20):
        family_keys = list(eligible.keys())
        random.shuffle(family_keys)

        selected = []
        for fk in family_keys:
            if len(selected) >= n_compounds:
                break
            c = random.choice(eligible[fk])
            if c not in selected:
                selected.append(c)

        # Sprinkle inorganic
        inorganics = theme_pools['inorganic']
        if inorganics and random.random() < 0.3 and len(selected) < n_compounds:
            ino = random.choice(inorganics)
            if ino not in selected:
                selected.append(ino)

        if len(selected) >= 3 and _validate_no_cross(selected):
            selected.sort(key=lambda c: c.best_pka())
            return selected

    return None


def _pick_exception_question(theme_pools, n_compounds=4):
    """Pick a question centered on a famous exception, themes not interleaved."""
    pools = theme_pools['exception']
    if not pools:
        return None

    for _ in range(20):
        ex_family = random.choice(list(pools.keys()))
        exception_compounds = pools[ex_family]
        selected = [random.choice(exception_compounds)]
        target_pka = selected[0].best_pka()

        all_others = []
        for family, members in theme_pools['functional_group'].items():
            all_others.extend(members)
        for family, members in theme_pools['substituent'].items():
            all_others.extend(members)

        nearby = [c for c in all_others if c.best_pka() is not None and
                  abs(c.best_pka() - target_pka) < 10 and c not in selected]
        random.shuffle(nearby)

        for c in nearby:
            if len(selected) >= n_compounds:
                break
            if c not in selected:
                selected.append(c)

        inorganics = theme_pools['inorganic']
        if inorganics and len(selected) < n_compounds:
            for ino in random.sample(inorganics, min(len(inorganics), n_compounds - len(selected))):
                if ino not in selected:
                    selected.append(ino)

        if len(selected) >= 3 and _validate_no_cross(selected):
            selected.sort(key=lambda c: c.best_pka())
            return selected

    return None


def generate_question_set(compounds, n_questions=5, sub_per_q=3, cpq=4):
    """Generate a full set of ranking questions with thematic constraints."""
    theme_pools = _build_theme_groups(compounds)

    # Print pool summary
    n_inorg = len(theme_pools['inorganic'])
    n_subst = sum(len(v) for v in theme_pools['substituent'].values())
    n_func = sum(len(v) for v in theme_pools['functional_group'].values())
    n_exc = sum(len(v) for v in theme_pools['exception'].values())
    n_conc = sum(len(v) for v in theme_pools['conjugate_acid'].values())
    print(f"  Theme pools: {n_inorg} inorganic, {n_subst} substituent, "
          f"{n_func} functional_group, {n_exc} exception, {n_conc} conjugate_acid")

    questions = []

    ALL_SUBS = [
        ('substituent', 1),
        ('substituent', 2),
        ('functional_group', 1),
        ('functional_group', 2),
        ('functional_group', 3),
        ('exception', 3),
    ]
    EASY_SUBS = [s for s in ALL_SUBS if s[1] == 1]

    for qi in range(n_questions):
        # Sub 1: difficulty must be 1, type random
        type1, diff1 = random.choice(EASY_SUBS)
        # Sub 2/3: pick 2 from remaining (distinct from sub 1)
        rest = [(t, d) for (t, d) in ALL_SUBS if (t, d) != (type1, diff1)]
        random.shuffle(rest)
        type2, diff2 = rest[0]
        type3, diff3 = rest[1]
        sub_types = [(type1, diff1), (type2, diff2), (type3, diff3)]

        sub_questions = []

        for q_type, difficulty in sub_types:
            group = None

            if q_type == 'substituent':
                group = _pick_substituent_question(theme_pools, cpq, difficulty)
            elif q_type == 'functional_group':
                group = _pick_functional_group_question(theme_pools, cpq, difficulty)
            elif q_type == 'exception':
                group = _pick_exception_question(theme_pools, cpq)

            if group is None or len(group) < 3:
                # Fallback
                group = select_compounds_for_ranking(compounds, cpq)

            # Deduplicate names
            seen_names = set()
            unique_group = []
            for c in group:
                base_name = re.sub(r'\(.*', '', c.name).strip()
                if base_name not in seen_names:
                    seen_names.add(base_name)
                    unique_group.append(c)
                if len(unique_group) >= cpq:
                    break

            unique_group.sort(key=lambda c: c.best_pka() if c.best_pka() is not None else 999)

            if len(unique_group) >= 3:
                sub_questions.append({
                    'compounds': unique_group,
                    'answer': [c.best_pka() for c in unique_group],
                    'names': [c.name for c in unique_group]
                })

        if len(sub_questions) >= sub_per_q:
            questions.append({
                'title': '酸性排序',
                'subs': sub_questions[:sub_per_q]
            })
        else:
            questions.append({
                'title': '酸性排序',
                'subs': _generate_fallback_subs(compounds, sub_per_q, cpq)
            })

    return questions


def _generate_fallback_subs(compounds, n_subs, cpq):
    """Generate fallback sub-questions from any available compounds."""
    subs = []
    for _ in range(n_subs):
        group = select_compounds_for_ranking(compounds, cpq)
        group.sort(key=lambda c: c.best_pka() if c.best_pka() is not None else 999)
        if len(group) >= 3:
            subs.append({
                'compounds': group,
                'answer': [c.best_pka() for c in group],
                'names': [c.name for c in group]
            })
    return subs


# ── CDXML Generation ────────────────────────────────────────────────

def smiles_to_cdxml(smiles, mol_id=1):
    """Convert SMILES to CDXML string (ChemDraw XML format)."""
    cleaned = clean_smiles(smiles)
    mol = Chem.MolFromSmiles(cleaned)
    if mol is None:
        return None

    try:
        AllChem.Compute2DCoords(mol)
    except Exception:
        return None

    conf = mol.GetConformer()

    # Build CDXML
    cdxml_lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<!DOCTYPE CDXML SYSTEM "http://www.cambridgesoft.com/xml/cdxml.dtd">',
        '<CDXML'
        ' BondLength="14.40"'
        ' LabelFont="Arial"'
        ' LabelSize="10"'
        ' LineWidth="0.6"'
        ' CaptionFont="Arial"'
        ' CaptionSize="10"'
        ' HashSpacing="2.5"'
        ' MarginWidth="1.5"'
        '>',
        '  <page>',
        '    <fragment>',
    ]

    # Normalize coordinates
    positions = []
    for atom in mol.GetAtoms():
        p = conf.GetAtomPosition(atom.GetIdx())
        positions.append((p.x, p.y))

    if not positions:
        return None

    # Center and scale
    xs = [p[0] for p in positions]
    ys = [p[1] for p in positions]
    cx = (max(xs) + min(xs)) / 2
    cy = (max(ys) + min(ys)) / 2

    scale = 30.0  # CDXML coordinate units per angstrom

    # Add nodes (atoms)
    for atom in mol.GetAtoms():
        idx = atom.GetIdx()
        p = conf.GetAtomPosition(idx)
        x = (p.x - cx) * scale
        y = -(p.y - cy) * scale  # Flip Y for CDXML

        element = atom.GetSymbol()
        atomic_num = atom.GetAtomicNum()
        charge = atom.GetFormalCharge()
        isotope = atom.GetIsotope()
        num_h = atom.GetNumExplicitHs()
        # Also include implicit H count
        total_h = atom.GetTotalNumHs()

        node_attrs = f'id="{idx + 1}" p="{x:.2f} {y:.2f}"'
        node_attrs += f' Element="{atomic_num}"'

        if total_h > 0:
            node_attrs += f' NumHydrogens="{total_h}"'
        if charge != 0:
            node_attrs += f' FormalCharge="{charge}"'
        if isotope != 0:
            node_attrs += f' Isotope="{isotope}"'

        cdxml_lines.append(f'      <n {node_attrs}/>')

    # Add bonds
    bond_id = 1
    for bond in mol.GetBonds():
        a1 = bond.GetBeginAtomIdx() + 1
        a2 = bond.GetEndAtomIdx() + 1

        order = bond.GetBondTypeAsDouble()
        if order == 1.0:
            bond_order = "1"
        elif order == 2.0:
            bond_order = "2"
        elif order == 3.0:
            bond_order = "3"
        elif order == 1.5:
            bond_order = "4"  # aromatic
        else:
            bond_order = "1"

        # Get bond display info
        display = "Solid"
        bond_dir = bond.GetBondDir()
        if bond_dir == Chem.BondDir.BEGINWEDGE:
            display = "WedgedHashBegin"
        elif bond_dir == Chem.BondDir.BEGINDASH:
            display = "Dashed"

        b_attrs = f'id="{bond_id}" B="{a1}" E="{a2}" Order="{bond_order}"'
        if display != "Solid":
            b_attrs += f' Display="{display}"'

        cdxml_lines.append(f'      <b {b_attrs}/>')
        bond_id += 1

    cdxml_lines.append('    </fragment>')
    cdxml_lines.append('  </page>')
    cdxml_lines.append('</CDXML>')

    return '\n'.join(cdxml_lines)


# ── Structure Image Generation (ACS 1996 style) ─────────────────────

def mol_to_acs_png(smiles, width=ACS_IMAGE_WIDTH, height=ACS_IMAGE_HEIGHT):
    """Draw molecule as PNG with ACS 1996 styling using RDKit."""
    cleaned = clean_smiles(smiles)
    mol = Chem.MolFromSmiles(cleaned)
    if mol is None:
        return None

    try:
        AllChem.Compute2DCoords(mol)
    except Exception:
        return None

    drawer = rdMolDraw2D.MolDraw2DCairo(width, height)
    opts = drawer.drawOptions()

    # ACS 1996 style settings
    opts.bondLineWidth = 2.5
    opts.fixedBondLength = ACS_BOND_LENGTH
    opts.fixedBondLength = 30
    opts.useBWAtomPalette()           # Black & white
    opts.clearBackground = True
    opts.addStereoAnnotation = True
    opts.bondLineWidth = 3.0
    opts.scaleBondWidth = True
    opts.multipleBondOffset = 0.18

    drawer.DrawMolecule(mol)
    drawer.FinishDrawing()

    return drawer.GetDrawingText()  # PNG bytes


# ── DOCX Generation ─────────────────────────────────────────────────

def _cn_font(run, size=Pt(10.5)):
    """Apply 宋体 + Times New Roman, black, given size."""
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = size
    # Set East-Asian font to 宋体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')


def _add_chinese_run(paragraph, text, bold=False, size=Pt(10.5)):
    """Add a run with 宋体+Times New Roman, black."""
    run = paragraph.add_run(text)
    _cn_font(run, size=size)
    run.bold = bold
    return run


def create_docx(questions, output_path, cdxml_dir=None):
    """Create a DOCX file with embedded CDXML structures and questions."""
    if not HAS_DOCX:
        print("ERROR: python-docx required")
        return

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25
    # East-Asian font for Normal style
    try:
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass

    FONT_SIZE = Pt(10.5)  # 五号

    # ── Title ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_chinese_run(title_p, 'CChO 酸性排序专项练习', bold=True, size=Pt(16))

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_chinese_run(sub_p, '（每题有3小题，将化合物按酸性由强到弱排序）', size=Pt(9))

    doc.add_paragraph()  # spacer

    # ── Questions ──
    for qi, question in enumerate(questions):
        # Question header: "X.  theme"
        q_label = f"{qi + 1}. "
        q_p = doc.add_paragraph()
        q_p.paragraph_format.space_before = Pt(12)
        q_p.paragraph_format.space_after = Pt(8)
        run_label = q_p.add_run(q_label)
        run_label.font.name = 'Times New Roman'
        run_label.font.size = Pt(12)
        run_label.bold = True
        run_label.font.color.rgb = RGBColor(0, 0, 0)
        _add_chinese_run(q_p, question['title'], bold=True, size=Pt(12))

        for si, sub in enumerate(question['subs']):
            compounds = sub['compounds']
            random.shuffle(compounds)

            # Sub-question label: "X-Y  "
            sub_label = f"{qi + 1}-{si + 1}  "
            sub_p = doc.add_paragraph()
            sub_p.paragraph_format.space_before = Pt(8)
            sub_p.paragraph_format.space_after = Pt(4)
            # Number part bold
            run_num = sub_p.add_run(sub_label)
            run_num.font.name = 'Times New Roman'
            run_num.font.size = FONT_SIZE
            run_num.bold = True
            run_num.font.color.rgb = RGBColor(0, 0, 0)
            # Text part not bold
            _add_chinese_run(sub_p, '将下列化合物按酸性由强到弱排列：', bold=False, size=FONT_SIZE)

            # Create a table for structures (1 row, n columns)
            n_cols = len(compounds)
            table = doc.add_table(rows=2, cols=n_cols)
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Set column widths
            col_width = Cm(16.0 / n_cols)
            for ci in range(n_cols):
                for row in table.rows:
                    row.cells[ci].width = col_width

            # Structure row
            for ci, comp in enumerate(compounds):
                cell = table.cell(0, ci)
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                img_bytes = mol_to_acs_png(comp.smiles, 300, 220)
                if img_bytes:
                    img_stream = io.BytesIO(img_bytes)
                    run = p.add_run()
                    try:
                        run.add_picture(img_stream, width=Cm(3.2))
                    except Exception:
                        run.text = comp.smiles
                else:
                    run = p.add_run(comp.smiles)
                    run.font.size = Pt(8)

                # Save CDXML
                if cdxml_dir:
                    cdxml = smiles_to_cdxml(comp.smiles)
                    if cdxml:
                        safe_name = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fff]', '_', comp.name)[:40]
                        cdxml_path = os.path.join(cdxml_dir, f'Q{qi+1}S{si+1}_{ci+1}_{safe_name}.cdxml')
                        try:
                            with open(cdxml_path, 'w', encoding='utf-8') as f:
                                f.write(cdxml)
                        except Exception:
                            pass

            # Label row — Arial font for letter labels
            for ci, comp in enumerate(compounds):
                cell = table.cell(1, ci)
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                label = chr(65 + ci)  # A, B, C, D, E
                run = p.add_run(label)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Answer blank
            answer_p = doc.add_paragraph()
            answer_p.paragraph_format.space_before = Pt(2)
            answer_p.paragraph_format.space_after = Pt(10)
            _add_chinese_run(answer_p, '酸性由强到弱：__________', size=FONT_SIZE)

        if qi < len(questions) - 1:
            doc.add_page_break()

    # ── Answer Key ──
    doc.add_page_break()
    ans_p = doc.add_paragraph()
    ans_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_chinese_run(ans_p, '参考答案', bold=True, size=Pt(16))

    for qi, question in enumerate(questions):
        ans_q_p = doc.add_paragraph()
        ans_q_p.paragraph_format.space_before = Pt(10)
        _add_chinese_run(ans_q_p, f"{qi + 1}. {question['title']}", bold=True, size=Pt(11))

        for si, sub in enumerate(question['subs']):
            comps = sub['compounds']
            labels = [chr(65 + i) for i in range(len(comps))]
            sorted_pairs = sorted(enumerate(comps), key=lambda x: x[1].best_pka())
            sorted_labels = ' > '.join([labels[i] for i, _ in sorted_pairs])
            names_str = ' > '.join([f'{c.name} (pKa={c.best_pka():.1f})' for _, c in sorted_pairs])

            p = doc.add_paragraph()
            _add_chinese_run(p, f'{qi + 1}-{si + 1}  {sorted_labels}   ', bold=False, size=FONT_SIZE)
            _add_chinese_run(p, names_str, size=FONT_SIZE)

    doc.save(output_path)
    print(f"DOCX saved to: {output_path}")
    return output_path


# ── Main ────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Generate CChO acid ranking questions')
    parser.add_argument('--count', type=int, default=1, help='Number of main questions')
    parser.add_argument('--output', type=str, help='Output DOCX path')
    parser.add_argument('--library', type=str, help='Path to 酸性数据.md')
    args = parser.parse_args()

    # Locate library
    library_path = args.library
    if not library_path:
        library_path = os.path.join(os.path.expanduser("~"), "Downloads", "酸性数据.md")

    print(f"Loading library...")
    compounds, stats = parse_acid_library(library_path)
    total = len(compounds)
    print(f"Total: {total} compounds ({stats['md']} from library + {stats['iupac']} from IUPAC)")

    if total < 15:
        print("ERROR: Not enough compounds in library (need at least 15)")
        sys.exit(1)

    # Generate questions
    questions = generate_question_set(compounds, n_questions=args.count)
    print(f"Generated {len(questions)} question groups")

    # Output path
    output_path = args.output
    if not output_path:
        output_path = os.path.join(os.path.expanduser("~"), "Downloads", "CChO酸性排序题.docx")
    ts = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    base, ext = os.path.splitext(output_path)
    output_path = f"{base}{ts}{ext}"

    # CDXML directory
    cdxml_dir = os.path.join(os.path.dirname(output_path), "cdxml_structures")
    os.makedirs(cdxml_dir, exist_ok=True)

    # Create DOCX
    create_docx(questions, output_path, cdxml_dir=cdxml_dir)

    print(f"\nDone! Generated {len(questions)} questions with {sum(len(q['subs']) for q in questions)} sub-questions.")
    print(f"Output: {output_path}")
    print(f"CDXML files: {cdxml_dir}/")


if __name__ == '__main__':
    main()

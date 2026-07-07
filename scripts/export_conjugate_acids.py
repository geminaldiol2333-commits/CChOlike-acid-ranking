#!/usr/bin/env python3
"""Export the ~28 conjugate_acid compounds with protonated CDXML into DOCX."""

import csv, os, sys, io, re
from collections import defaultdict
from datetime import datetime

from rdkit import Chem, RDLogger
from rdkit.Chem import AllChem, Draw
from rdkit.Chem.Draw import rdMolDraw2D
RDLogger.logger().setLevel(RDLogger.ERROR)

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CSV_PATH = os.path.join(os.path.expanduser("~"), "Downloads",
                        "Dissociation-Constants", "iupac_high-confidence_v2_3.csv")
PKA_TYPE_RANK = {'pKa1': 0, 'pKa': 1, 'pKaH1': 2,
                 'pKa2': 3, 'pKa3': 4, 'pKa4': 5, 'pKa5': 6}
ASSESSMENT_RANK = {'Reliable': 0, 'Approximate': 1, 'Uncertain': 2}
TEMP_RANK = {'25': 0, '20': 1, '': 2}
VALID_ACIDITY_LABELS = {'AH', 'A', ''}
PKA_MIN, PKA_MAX = -20, 60


def load_problematic(csv_path):
    """Load only pKaH1 entries that survive dedup, have neutral SMILES,
    AND would have failed the old N-only protonation (no N to protonate)."""
    entries_by_smiles = {}
    with open(csv_path, 'r', encoding='utf-8', errors='replace') as f:
        reader = csv.reader(f)
        next(reader, None)
        for row in reader:
            if len(row) < 21: continue
            smiles = row[1].strip()
            pka_type = row[3].strip()
            pka_str = row[4].strip()
            T = row[5].strip()
            assessment = row[8].strip()
            acidity_label = row[18].strip()
            cosolvent = row[20].strip()
            name = row[12].strip()
            if assessment not in ASSESSMENT_RANK: continue
            if cosolvent: continue
            if T not in TEMP_RANK: continue
            if pka_type not in PKA_TYPE_RANK: continue
            if acidity_label not in VALID_ACIDITY_LABELS: continue
            try: pka = float(pka_str)
            except ValueError: continue
            if pka < PKA_MIN or pka > PKA_MAX: continue
            if not smiles: continue
            if not pka_type.startswith('pKaH'): continue
            if '+' in smiles: continue
            sort_key = (TEMP_RANK[T], ASSESSMENT_RANK[assessment],
                        PKA_TYPE_RANK[pka_type], pka)
            if smiles not in entries_by_smiles:
                entries_by_smiles[smiles] = []
            entries_by_smiles[smiles].append(
                {'smiles': smiles, 'pka': pka, 'pka_type': pka_type,
                 'name': name, 'assessment': assessment, 'sort_key': sort_key})

    result = []
    for smiles, entries in entries_by_smiles.items():
        entries.sort(key=lambda e: e['sort_key'])
        best = entries[0]
        if best['pka_type'].startswith('pKaH'):
            result.append(best)

    # Filter: keep only those that would FAIL N-only protonation
    # (no protonatable N, i.e. old function would have returned original)
    def has_protonatable_N(smi):
        if '+' in smi: return True
        try:
            mol = Chem.MolFromSmiles(smi)
            if mol is None: return True  # can't parse → keep
            for atom in mol.GetAtoms():
                if atom.GetAtomicNum() == 7 and atom.GetFormalCharge() == 0:
                    heavy = sum(1 for b in atom.GetBonds()
                               if b.GetOtherAtom(atom).GetAtomicNum() != 1)
                    if heavy < 4:
                        return True
        except Exception:
            return True
        return False

    result = [r for r in result if not has_protonatable_N(r['smiles'])]
    return result


def protonate_any(smiles):
    """Protonate most basic site: N > O > S > P."""
    if '+' in smiles:
        return smiles, False
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        return smiles, False
    max_valence = {7: 4, 8: 3, 16: 3, 15: 4}
    candidates = []
    for atom in mol.GetAtoms():
        anum = atom.GetAtomicNum()
        if anum not in max_valence: continue
        if atom.GetFormalCharge() != 0: continue
        heavy_bonds = sum(1 for b in atom.GetBonds()
                          if b.GetOtherAtom(atom).GetAtomicNum() != 1)
        if heavy_bonds >= max_valence[anum]: continue
        n_h = atom.GetTotalNumHs()
        score = -heavy_bonds + n_h * 2
        elem_prio = {7: 0, 8: 1, 16: 1, 15: 2}[anum]
        candidates.append((score, -elem_prio, atom.GetIdx()))
    if not candidates:
        return smiles, False
    candidates.sort(reverse=True)
    _, _, best_idx = candidates[0]
    try:
        mol_h = Chem.RWMol(mol)
        atom = mol_h.GetAtomWithIdx(best_idx)
        atom.SetNumExplicitHs(atom.GetNumExplicitHs() + 1)
        atom.SetFormalCharge(1)
        mol_h = mol_h.GetMol()
        Chem.SanitizeMol(mol_h)
        new_smiles = Chem.MolToSmiles(mol_h)
        if '+' in new_smiles:
            return new_smiles, True
        return smiles, False
    except Exception:
        return smiles, False


def mol_to_png(smiles, w=350, h=260):
    mol = Chem.MolFromSmiles(smiles)
    if mol is None: return None
    try: AllChem.Compute2DCoords(mol)
    except Exception: return None
    drawer = rdMolDraw2D.MolDraw2DCairo(w, h)
    opts = drawer.drawOptions()
    opts.bondLineWidth = 2.5
    opts.fixedBondLength = 26
    opts.useBWAtomPalette()
    opts.clearBackground = True
    opts.scaleBondWidth = True
    drawer.DrawMolecule(mol)
    drawer.FinishDrawing()
    return drawer.GetDrawingText()


def smiles_to_cdxml(smiles):
    mol = Chem.MolFromSmiles(smiles)
    if mol is None: return None
    try: AllChem.Compute2DCoords(mol)
    except Exception: return None
    conf = mol.GetConformer()
    positions = [(conf.GetAtomPosition(a.GetIdx()).x,
                  conf.GetAtomPosition(a.GetIdx()).y) for a in mol.GetAtoms()]
    if not positions: return None
    xs, ys = [p[0] for p in positions], [p[1] for p in positions]
    cx, cy = (max(xs)+min(xs))/2, (max(ys)+min(ys))/2
    scale = 30.0
    lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<!DOCTYPE CDXML SYSTEM "http://www.cambridgesoft.com/xml/cdxml.dtd">',
        '<CDXML BondLength="14.40" LabelFont="Arial" LabelSize="10" LineWidth="0.6">',
        '  <page><fragment>']
    for atom in mol.GetAtoms():
        p = conf.GetAtomPosition(atom.GetIdx())
        x, y = (p.x-cx)*scale, -(p.y-cy)*scale
        anum, charge, th = atom.GetAtomicNum(), atom.GetFormalCharge(), atom.GetTotalNumHs()
        a = f'id="{atom.GetIdx()+1}" p="{x:.2f} {y:.2f}" Element="{anum}"'
        if th > 0: a += f' NumHydrogens="{th}"'
        if charge != 0: a += f' FormalCharge="{charge}"'
        lines.append(f'      <n {a}/>')
    bid = 1
    for bond in mol.GetBonds():
        a1, a2 = bond.GetBeginAtomIdx()+1, bond.GetEndAtomIdx()+1
        o = {1.0:"1",2.0:"2",3.0:"3"}.get(bond.GetBondTypeAsDouble(),"1")
        d = "Solid"
        bd = bond.GetBondDir()
        if bd == Chem.BondDir.BEGINWEDGE: d = "WedgedHashBegin"
        elif bd == Chem.BondDir.BEGINDASH: d = "Dashed"
        ba = f'id="{bid}" B="{a1}" E="{a2}" Order="{o}"'
        if d != "Solid": ba += f' Display="{d}"'
        lines.append(f'      <b {ba}/>')
        bid += 1
    lines.append('    </fragment></page>')
    lines.append('</CDXML>')
    return '\n'.join(lines)


def _cn_font(run, size=Pt(9)):
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

def _add_zh(p, text, bold=False, size=Pt(9)):
    r = p.add_run(text); _cn_font(r, size); r.bold = bold; return r


def main():
    print("Loading problematic pKaH1 entries...")
    compounds = load_problematic(CSV_PATH)
    print(f"  Neutral pKaH1 surviving dedup: {len(compounds)}")

    # Protonate, keep all
    data = []
    n_ok = 0
    for c in compounds:
        ps, ok = protonate_any(c['smiles'])
        if ok: n_ok += 1
        data.append({
            'name': c['name'], 'pka': c['pka'], 'assessment': c['assessment'],
            'neutral_smiles': c['smiles'], 'protonated_smiles': ps, 'success': ok,
        })
    print(f"  Protonated: {n_ok} / {len(data)}")

    # Output path
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    out = os.path.join(os.path.expanduser("~"), "Downloads",
                       f"共轭酸质子化对照表_{ts}.docx")

    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(29.7); s.page_height = Cm(21.0)
        s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_zh(tp, '共轭酸质子化对照表', bold=True, size=Pt(14))

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_zh(sp, f'共 {len(data)} 个 pKaH1 条目（成功质子化 {n_ok} 个）', size=Pt(8))
    doc.add_paragraph()

    # CDXML output dir
    cdxml_dir = os.path.join(os.path.expanduser("~"), "Downloads", "conjugate_acid_cdxml")
    os.makedirs(cdxml_dir, exist_ok=True)

    for i, cd in enumerate(data, 1):
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(10)
        _add_zh(hp, f'{i}. {cd["name"][:90]}', bold=True, size=Pt(9))
        tag = ' (已质子化)' if cd['success'] else ' (质子化失败)'
        _add_zh(hp, f'  pKaH1={cd["pka"]:.1f}  {cd["assessment"]}{tag}', size=Pt(8))

        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ci, t in enumerate(['原始中性形式', '质子化共轭酸']):
            c = table.cell(0, ci); c.paragraphs[0].clear()
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_zh(p, t, bold=True, size=Pt(8))
        for row in table.rows:
            row.cells[0].width = Cm(12.5); row.cells[1].width = Cm(12.5)

        # Neutral
        nimg = mol_to_png(cd['neutral_smiles'])
        if nimg:
            c = table.cell(1, 0); c.paragraphs[0].clear()
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(nimg), width=Cm(4.5))
            cp = c.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_zh(cp, cd['neutral_smiles'], size=Pt(6))

        # Protonated
        pimg = mol_to_png(cd['protonated_smiles'])
        if pimg:
            c = table.cell(1, 1); c.paragraphs[0].clear()
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(pimg), width=Cm(4.5))
            cp = c.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_zh(cp, cd['protonated_smiles'], size=Pt(6))

        # Export CDXML
        safe = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fff]', '_', cd['name'])[:35]
        for suffix, smi in [('neutral', cd['neutral_smiles']), ('protonated', cd['protonated_smiles'])]:
            cdxml = smiles_to_cdxml(smi)
            if cdxml:
                with open(os.path.join(cdxml_dir, f'{i:03d}_{suffix}_{safe}.cdxml'), 'w', encoding='utf-8') as f:
                    f.write(cdxml)

        if i % 20 == 0:
            print(f"  Progress: {i}/{len(data)}")

    doc.save(out)
    print(f"\nDOCX saved: {out}")
    print(f"CDXML exported: {cdxml_dir}/")


if __name__ == '__main__':
    main()
